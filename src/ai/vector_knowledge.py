"""
向量知识库服务 - 使用OpenAI Embeddings实现语义搜索

功能：
1. 文档上传和解析（支持PDF、Word、TXT等）
2. 文本分块和向量化
3. 向量检索（语义搜索）
4. 知识库管理
"""

import os
import json
import hashlib
import numpy as np
from typing import List, Dict, Optional, Tuple
from datetime import datetime
import asyncio

# 文档解析库
import PyPDF2
import docx
from io import BytesIO

# OpenAI
from openai import AsyncOpenAI


class VectorKnowledgeService:
    """向量知识库服务"""
    
    def __init__(self):
        # 使用系统配置的 AI Hub Mix API
        self.api_key = os.getenv('AIHUBMIX_API_KEY', 'sk-5dn0RF7nn31mpHNjEfC5Ca1579F447418aE48e7b0d8b18F7')
        self.client = AsyncOpenAI(
            api_key=self.api_key,
            base_url=os.getenv('AIHUBMIX_BASE_URL', 'https://aihubmix.com/v1')
        )
        self.embedding_model = "text-embedding-3-small"  # 更便宜的embedding模型
    
    async def create_embedding(self, text: str) -> List[float]:
        """
        创建文本向量
        
        参数:
            text: 要向量化的文本
            
        返回:
            向量数组
        """
        try:
            response = await self.client.embeddings.create(
                input=text,
                model=self.embedding_model
            )
            return response.data[0].embedding
        except Exception as e:
            print(f"❌ 创建向量失败: {str(e)}")
            raise
    
    async def batch_create_embeddings(self, texts: List[str]) -> List[List[float]]:
        """
        批量创建文本向量
        
        参数:
            texts: 文本列表
            
        返回:
            向量列表
        """
        try:
            response = await self.client.embeddings.create(
                input=texts,
                model=self.embedding_model
            )
            return [item.embedding for item in response.data]
        except Exception as e:
            print(f"❌ 批量创建向量失败: {str(e)}")
            raise
    
    def parse_pdf(self, file_content: bytes) -> str:
        """解析PDF文件"""
        try:
            pdf_file = BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            
            return text.strip()
        except Exception as e:
            print(f"❌ 解析PDF失败: {str(e)}")
            raise
    
    def parse_docx(self, file_content: bytes) -> str:
        """解析Word文档"""
        try:
            doc_file = BytesIO(file_content)
            doc = docx.Document(doc_file)
            
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            return text.strip()
        except Exception as e:
            print(f"❌ 解析Word文档失败: {str(e)}")
            raise
    
    def parse_txt(self, file_content: bytes) -> str:
        """解析纯文本文件"""
        try:
            # 尝试多种编码
            for encoding in ['utf-8', 'gbk', 'gb2312', 'latin-1']:
                try:
                    return file_content.decode(encoding)
                except UnicodeDecodeError:
                    continue
            raise ValueError("无法识别文件编码")
        except Exception as e:
            print(f"❌ 解析文本文件失败: {str(e)}")
            raise
    
    def parse_document(self, file_content: bytes, filename: str) -> str:
        """
        解析文档
        
        参数:
            file_content: 文件内容（字节）
            filename: 文件名
            
        返回:
            解析后的文本
        """
        extension = filename.lower().split('.')[-1]
        
        if extension == 'pdf':
            return self.parse_pdf(file_content)
        elif extension in ['doc', 'docx']:
            return self.parse_docx(file_content)
        elif extension == 'txt':
            return self.parse_txt(file_content)
        else:
            raise ValueError(f"不支持的文件类型: {extension}")
    
    def split_text(
        self, 
        text: str, 
        chunk_size: int = 500, 
        overlap: int = 50
    ) -> List[str]:
        """
        文本分块
        
        参数:
            text: 原始文本
            chunk_size: 每块大小（字符数）
            overlap: 重叠字符数
            
        返回:
            文本块列表
        """
        chunks = []
        start = 0
        text_length = len(text)
        
        while start < text_length:
            end = start + chunk_size
            chunk = text[start:end]
            
            # 如果不是最后一块，尝试在句号处截断
            if end < text_length:
                last_period = chunk.rfind('.')
                last_newline = chunk.rfind('\n')
                last_break = max(last_period, last_newline)
                
                if last_break > chunk_size * 0.5:  # 至少保留一半内容
                    end = start + last_break + 1
                    chunk = text[start:end]
            
            chunks.append(chunk.strip())
            start = end - overlap
        
        return [c for c in chunks if c]  # 过滤空块
    
    def calculate_file_hash(self, file_content: bytes) -> str:
        """计算文件哈希值（用于去重）"""
        return hashlib.md5(file_content).hexdigest()
    
    def cosine_similarity(self, vec1: List[float], vec2: List[float]) -> float:
        """
        计算余弦相似度
        
        参数:
            vec1, vec2: 向量数组
            
        返回:
            相似度 (0-1)
        """
        vec1_arr = np.array(vec1)
        vec2_arr = np.array(vec2)
        
        dot_product = np.dot(vec1_arr, vec2_arr)
        norm1 = np.linalg.norm(vec1_arr)
        norm2 = np.linalg.norm(vec2_arr)
        
        if norm1 == 0 or norm2 == 0:
            return 0.0
        
        return float(dot_product / (norm1 * norm2))
    
    async def upload_document(
        self,
        file_content: bytes,
        filename: str,
        title: str,
        category: str = "general",
        description: str = None,
        db_session = None
    ) -> Dict:
        """
        上传文档到知识库
        
        参数:
            file_content: 文件内容（字节）
            filename: 文件名
            title: 文档标题
            category: 分类
            description: 描述
            db_session: 数据库会话
            
        返回:
            文档信息
        """
        from src.crm.database import get_session, KnowledgeDocument, KnowledgeChunk
        
        if db_session is None:
            db_session = get_session()
        
        try:
            # 0. 检查文件是否已存在
            file_hash = self.calculate_file_hash(file_content)
            existing_doc = db_session.query(KnowledgeDocument).filter(
                KnowledgeDocument.file_hash == file_hash,
                KnowledgeDocument.is_active == True
            ).first()
            
            if existing_doc:
                raise ValueError(f"该文件已存在于知识库中：'{existing_doc.title}'（文件名：{existing_doc.filename}）")
            
            # 1. 解析文档
            print(f"📝 解析文档: {filename}")
            text = self.parse_document(file_content, filename)
            
            # 2. 文本分块
            print(f"✂️ 分块文本...")
            chunks = self.split_text(text)
            print(f"✅ 生成 {len(chunks)} 个分块")
            
            # 3. 创建文档记录
            document = KnowledgeDocument(
                title=title,
                filename=filename,
                category=category,
                summary=description,  # 使用summary字段
                file_size=len(file_content),
                file_hash=file_hash,
                chunk_count=len(chunks),
                status='completed',
                content=text[:5000],  # 保存前5000字符作为预览
                created_at=datetime.utcnow()
            )
            db_session.add(document)
            db_session.flush()
            
            # 4. 向量化所有分块
            print(f"🧪 向量化文本...")
            chunk_texts = [chunk for chunk in chunks]
            embeddings = await self.batch_create_embeddings(chunk_texts)
            
            # 5. 保存分块和向量
            print(f"💾 保存到数据库...")
            for idx, (chunk_text, embedding) in enumerate(zip(chunks, embeddings)):
                chunk = KnowledgeChunk(
                    document_id=document.id,
                    content=chunk_text,
                    chunk_index=idx,
                    embedding=json.dumps(embedding),  # 以JSON格式存储
                    chunk_metadata=json.dumps({}),
                    token_count=len(chunk_text) // 4,  # 粗略估计
                    char_count=len(chunk_text),
                    created_at=datetime.utcnow()
                )
                db_session.add(chunk)
            
            db_session.commit()
            
            print(f"✅ 文档上传成功: {title}")
            
            return {
                "id": document.id,
                "title": document.title,
                "filename": document.filename,
                "category": document.category,
                "chunk_count": len(chunks)
            }
            
        except Exception as e:
            db_session.rollback()
            print(f"❌ 文档上传失败: {str(e)}")
            raise
        finally:
            if db_session:
                db_session.close()
    
    async def search_similar(
        self,
        query: str,
        limit: int = 5,
        category: Optional[str] = None,
        min_similarity: float = 0.3,  # 🔥 新增：最低相似度阈值
        db_session = None
    ) -> List[Dict]:
        """
        向量相似度搜索
        
        参数:
            query: 查询文本
            limit: 返回结果数量
            category: 知识库分类（可选）
            min_similarity: 最低相似度阈值（默认0.3，过滤低相关内容）
            db_session: 数据库会话
            
        返回:
            相似文档列表
        """
        from src.crm.database import get_session, KnowledgeChunk, KnowledgeDocument
        
        if db_session is None:
            db_session = get_session()
        
        try:
            # 1. 将查询文本向量化
            query_vector = await self.create_embedding(query)
            
            # 2. 从数据库获取所有活跃的分块
            if category:
                chunks = db_session.query(KnowledgeChunk).join(
                    KnowledgeDocument
                ).filter(
                    KnowledgeChunk.is_active == True,
                    KnowledgeDocument.category == category
                ).all()
            else:
                chunks = db_session.query(KnowledgeChunk).filter(
                    KnowledgeChunk.is_active == True
                ).all()
            
            # 3. 计算每个分块的相似度
            results = []
            for chunk in chunks:
                if not chunk.embedding:
                    continue
                
                try:
                    chunk_vector = json.loads(chunk.embedding)
                    similarity = self.cosine_similarity(query_vector, chunk_vector)
                    
                    # 🔥 新增：过滤低于阈值的结果
                    if similarity < min_similarity:
                        continue
                    
                    results.append({
                        "id": chunk.id,
                        "document_id": chunk.document_id,
                        "document_title": chunk.document.title if chunk.document else "Unknown",
                        "content": chunk.content,
                        "chunk_index": chunk.chunk_index,
                        "metadata": json.loads(chunk.chunk_metadata) if chunk.chunk_metadata else {},
                        "similarity": similarity
                    })
                except Exception as e:
                    print(f"⚠️ 跳过无效分块 {chunk.id}: {str(e)}")
                    continue
            
            # 4. 按相似度排序并返回前 N 个
            results.sort(key=lambda x: x["similarity"], reverse=True)
            
            # 🔥 新增：记录日志
            if results:
                print(f"✅ 找到 {len(results)} 个相关知识片段（相似度 >= {min_similarity})）")
                print(f"   最高相似度: {results[0]['similarity']:.2f}")
            else:
                print(f"⚠️ 未找到相似度 >= {min_similarity} 的知识片段")
            
            return results[:limit]
            
        except Exception as e:
            print(f"❌ 向量搜索失败: {str(e)}")
            raise
        finally:
            if db_session:
                db_session.close()


# 全局实例
_vector_service = None

def get_vector_service() -> VectorKnowledgeService:
    """获取向量知识库服务单例"""
    global _vector_service
    if _vector_service is None:
        _vector_service = VectorKnowledgeService()
    return _vector_service
